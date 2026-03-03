"""
Resume Parser Module

This module handles parsing of resume files in different formats (PDF, DOCX, TXT)
and extraction of text content.
"""

import os
import logging
from typing import Tuple, Optional
from pathlib import Path

logger = logging.getLogger(__name__)


class ResumeParser:
    """
    Parse resume files in multiple formats (PDF, DOCX, TXT).
    
    Supported Formats:
        - PDF (.pdf)
        - DOCX (.docx)
        - Text (.txt)
    """
    
    SUPPORTED_FORMATS = ['.pdf', '.docx', '.txt', '.doc']
    
    def __init__(self):
        """Initialize Resume Parser."""
        self._check_dependencies()
        logger.info("Resume Parser initialized")
    
    def _check_dependencies(self):
        """Check if required dependencies are available."""
        try:
            import pdfplumber
            self.pdfplumber = pdfplumber
        except ImportError:
            logger.warning("pdfplumber not installed. PDF parsing will not work.")
            self.pdfplumber = None
        
        try:
            import docx
            self.docx = docx
        except ImportError:
            logger.warning("python-docx not installed. DOCX parsing will not work.")
            self.docx = None
    
    def is_supported_format(self, file_path: str) -> bool:
        """
        Check if file format is supported.
        
        Args:
            file_path (str): Path to file
            
        Returns:
            bool: True if format is supported
        """
        file_ext = Path(file_path).suffix.lower()
        return file_ext in self.SUPPORTED_FORMATS
    
    def parse_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path (str): Path to PDF file
            
        Returns:
            str: Extracted text
            
        Raises:
            ImportError: If pdfplumber is not installed
            FileNotFoundError: If file doesn't exist
        """
        if not self.pdfplumber:
            raise ImportError("pdfplumber is required for PDF parsing. "
                            "Install it using: pip install pdfplumber")
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        text = ""
        try:
            with self.pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            logger.info(f"Successfully extracted text from PDF: {file_path}")
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            raise
        
        return text.strip()
    
    def parse_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path (str): Path to DOCX file
            
        Returns:
            str: Extracted text
            
        Raises:
            ImportError: If python-docx is not installed
            FileNotFoundError: If file doesn't exist
        """
        if not self.docx:
            raise ImportError("python-docx is required for DOCX parsing. "
                            "Install it using: pip install python-docx")
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        text = ""
        try:
            doc = self.docx.Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            logger.info(f"Successfully extracted text from DOCX: {file_path}")
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise
        
        return text.strip()
    
    def parse_txt(self, file_path: str, encoding: str = 'utf-8') -> str:
        """
        Extract text from TXT file.
        
        Args:
            file_path (str): Path to TXT file
            encoding (str): File encoding
            
        Returns:
            str: File content
            
        Raises:
            FileNotFoundError: If file doesn't exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                text = file.read()
            
            logger.info(f"Successfully extracted text from TXT: {file_path}")
            return text.strip()
        except Exception as e:
            logger.error(f"Error parsing TXT {file_path}: {e}")
            raise
    
    def parse_file(self, file_path: str) -> str:
        """
        Parse resume file based on its format.
        
        Args:
            file_path (str): Path to resume file
            
        Returns:
            str: Extracted text content
            
        Raises:
            ValueError: If format is not supported
            FileNotFoundError: If file doesn't exist
        """
        file_ext = Path(file_path).suffix.lower()
        
        if not self.is_supported_format(file_path):
            raise ValueError(f"Unsupported file format: {file_ext}. "
                           f"Supported formats: {', '.join(self.SUPPORTED_FORMATS)}")
        
        # Route to appropriate parser
        if file_ext == '.pdf':
            return self.parse_pdf(file_path)
        elif file_ext == '.docx':
            return self.parse_docx(file_path)
        elif file_ext in ['.txt', '.doc']:
            return self.parse_txt(file_path)
        else:
            raise ValueError(f"Unable to parse format: {file_ext}")
    
    def batch_parse(self, file_paths: list) -> dict:
        """
        Parse multiple resume files.
        
        Args:
            file_paths (list): List of file paths
            
        Returns:
            dict: Dictionary with file paths as keys and extracted text as values
        """
        results = {}
        
        for file_path in file_paths:
            try:
                text = self.parse_file(file_path)
                results[file_path] = text
            except Exception as e:
                logger.error(f"Failed to parse {file_path}: {e}")
                results[file_path] = None
        
        logger.info(f"Batch parsed {len([v for v in results.values() if v])} "
                   f"out of {len(file_paths)} files successfully")
        
        return results
